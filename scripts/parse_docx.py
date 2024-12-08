from docx import Document
import json

def parse_docx(file_path):
    """
    Парсинг таблиц и списков из docx-файла.

    :param file_path: Путь к docx-файлу
    :return: Данные в виде словаря
    """
    doc = Document(file_path)
    parsed_data = {"tables": [], "lists": []}

    # Парсинг таблиц
    for table in doc.tables:
        table_data = [[cell.text.strip() for cell in row.cells] for row in table.rows]
        parsed_data["tables"].append(table_data)

    # Парсинг списков
    for paragraph in doc.paragraphs:
        if paragraph.style.name.startswith("List"):
            parsed_data["lists"].append(paragraph.text.strip())

    return parsed_data

if __name__ == "__main__":
    file_path = "data/sample.docx"
    parsed_data = parse_docx(file_path)

    # Сохранение в JSON
    with open("data/parsed_data.json", "w", encoding="utf-8") as f:
        json.dump(parsed_data, f, ensure_ascii=False, indent=4)

    print("Данные успешно извлечены и сохранены.")