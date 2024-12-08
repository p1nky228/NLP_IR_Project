from docx import Document


def create_simple_docx(file_path):
    doc = Document()
    doc.add_heading('Пример документа', level=1)

    # Добавление таблицы
    table = doc.add_table(rows=2, cols=3)
    table.style = 'Table Grid'
    headers = ["Header 1", "Header 2", "Header 3"]
    values = [["Value 1", "Value 2", "Value 3"]]

    for i, cell in enumerate(table.rows[0].cells):
        cell.text = headers[i]
    for i, cell in enumerate(table.rows[1].cells):
        cell.text = values[0][i]

    # Добавление списка
    doc.add_paragraph("Пример списка:")
    doc.add_paragraph("Первый элемент", style="List Bullet")
    doc.add_paragraph("Второй элемент", style="List Bullet")
    doc.add_paragraph("Подэлемент", style="List Number 2")

    # Сохранение файла
    doc.save(file_path)





def create_hard_docx(file_path):
    # Создаем сложный документ
    doc = Document()

    # Добавляем заголовок
    doc.add_heading('Сложный тестовый документ', level=1)

    # Добавляем подразделы
    doc.add_heading('Раздел 1: Общая информация', level=2)
    doc.add_paragraph(
        "Этот документ предназначен для тестирования программы парсинга. "
        "Он содержит различные типы данных, включая таблицы, списки, изображения и стилизованный текст."
    )

    # Таблица с данными
    doc.add_heading('Пример таблицы', level=2)
    table = doc.add_table(rows=3, cols=3)
    table.style = 'Light Grid Accent 1'

    # Заполняем таблицу
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Колонка 1'
    header_cells[1].text = 'Колонка 2'
    header_cells[2].text = 'Колонка 3'

    data = [
        ['Данные 1.1', 'Данные 1.2', 'Данные 1.3'],
        ['Данные 2.1', 'Данные 2.2', 'Данные 2.3']
    ]

    for i, row in enumerate(table.rows[1:]):
        for j, cell in enumerate(row.cells):
            cell.text = data[i][j]

    # Списки
    doc.add_heading('Пример списка', level=2)
    doc.add_paragraph('Это маркированный список:', style='Intense Quote')
    doc.add_paragraph('Пункт 1', style='List Bullet')
    doc.add_paragraph('Пункт 2', style='List Bullet')
    doc.add_paragraph('Пункт 3', style='List Bullet')

    doc.add_paragraph('Это нумерованный список:', style='Intense Quote')
    doc.add_paragraph('Первый пункт', style='List Number')
    doc.add_paragraph('Второй пункт', style='List Number')
    doc.add_paragraph('Третий пункт', style='List Number')

    # Сложное форматирование текста
    doc.add_heading('Пример форматирования текста', level=2)
    paragraph = doc.add_paragraph()
    paragraph.add_run('Этот текст жирный. ').bold = True
    paragraph.add_run('Этот текст курсивный. ').italic = True
    paragraph.add_run('Этот текст подчеркнутый. ').underline = True
    paragraph.add_run('Этот текст выделен желтым. ').font.highlight_color = 5

    # Сохранение файла
    doc.save(file_path)

if __name__ == "__main__":
    create_simple_docx("data/simple.docx")
    #create_hard_docx("data/hard.docx")
    print("Пример документа создан.")